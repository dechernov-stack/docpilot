from __future__ import annotations

import difflib
import hashlib
import html
import json
import logging
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import uuid
from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from .database import Database
from .sdoc import fingerprint, generate_rendered_sdoc, generate_structure_sdoc


NUMBER_RE = re.compile(
    r"(?<![\w-])(?:[≤≥]\s*)?-?\d+(?:[.,]\d+)?(?:\s*(?:[%°]?[A-Za-zА-Яа-яЁё³]+(?:/[A-Za-zА-Яа-яЁё³]+)*))?",
    re.IGNORECASE,
)
MARKER_RE = re.compile(r"\[([A-Z]{1,12}-\d{3})\]")
QUALIFIER_RE = re.compile(
    r"(?<!\w)(?:не\s+(?:менее|более|допуска\w*|долж\w*|мож\w*|разреш\w*)|"
    r"не|без|только|кроме|запрещ\w*|обязат\w*)(?!\w)",
    re.IGNORECASE,
)
logger = logging.getLogger("docpilot")


def find_tool(name: str) -> str | None:
    configured = os.getenv(f"{name.upper()}_BIN")
    candidates = [configured, shutil.which(name), str(Path(sys.executable).with_name(name))]
    return next((candidate for candidate in candidates if candidate and Path(candidate).is_file()), None)


class DocPilotError(Exception):
    def __init__(self, status: int, code: str, message: str, reasons: Iterable[str] = ()):
        super().__init__(message)
        self.status = status
        self.code = code
        self.message = message
        self.reasons = list(reasons)


def utcnow() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def _json(value: str | None, default: Any) -> Any:
    if not value:
        return default
    return json.loads(value)


def _dict(row: sqlite3.Row | None) -> dict[str, Any] | None:
    return dict(row) if row is not None else None


def _normalize_number(token: str) -> str:
    match = re.search(r"-?\d+(?:[.,]\d+)?", token)
    if not match:
        return token
    raw = match.group(0).replace(",", ".").lstrip("+")
    try:
        number = f"{float(raw):g}"
    except ValueError:
        number = raw
    operator = "≤" if "≤" in token[:match.start()] else "≥" if "≥" in token[:match.start()] else ""
    original_unit = re.sub(r"\s+", " ", token[match.end():].strip())
    unit = original_unit.lower()
    if unit.startswith("сек") or unit == "с":
        unit = "с"
    elif unit.startswith("станц"):
        unit = "станция"
    elif unit.startswith("оператор"):
        unit = "оператор"
    elif unit.startswith("месяц"):
        unit = "месяц"
    elif original_unit == "В":
        unit = "В"
    elif unit == "°c":
        unit = "°C"
    quantity = f"{number} {unit}" if unit else number
    return f"{operator}{quantity}" if operator else quantity


def extract_numbers(text: str) -> list[str]:
    scrubbed = MARKER_RE.sub("", text)
    scrubbed = re.sub(r"\b\d{4}-\d{2}-\d{2}(?:T[^\s]+)?\b", "", scrubbed)
    scrubbed = re.sub(r"§\s*\d+(?:\.\d+)*", "", scrubbed)
    return [_normalize_number(match.group(0)) for match in NUMBER_RE.finditer(scrubbed)]


def extract_qualifiers(text: str) -> list[str]:
    return sorted(re.sub(r"\s+", " ", match.group(0).lower()) for match in QUALIFIER_RE.finditer(text))


def _display_number(value: float | int | None) -> str:
    if value is None:
        return ""
    numeric = float(value)
    return str(int(numeric)) if numeric.is_integer() else str(numeric).replace(".", ",")


class GitRepository:
    """GitPython adapter with a CLI fallback; no global identity is used."""

    def __init__(self, path: Path):
        self.path = Path(path)

    def _run(self, *args: str, actor: str = "system", check: bool = True) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        env.update(
            {
                "GIT_AUTHOR_NAME": actor,
                "GIT_AUTHOR_EMAIL": f"{actor}@docpilot.local",
                "GIT_COMMITTER_NAME": actor,
                "GIT_COMMITTER_EMAIL": f"{actor}@docpilot.local",
            }
        )
        result = subprocess.run(
            ["git", *args],
            cwd=self.path,
            env=env,
            text=True,
            capture_output=True,
            check=False,
        )
        if check and result.returncode != 0:
            raise DocPilotError(500, "GIT_ERROR", "Операция Git не выполнена", [result.stderr.strip() or result.stdout.strip()])
        return result

    def initialize(self):
        self.path.mkdir(parents=True, exist_ok=True)
        try:
            from git import Repo

            if (self.path / ".git").exists():
                return Repo(self.path)
            try:
                return Repo.init(self.path, initial_branch="main")
            except TypeError:
                return Repo.init(self.path)
        except ImportError:
            if not (self.path / ".git").exists():
                result = subprocess.run(["git", "init", "-b", "main"], cwd=self.path, text=True, capture_output=True, check=False)
                if result.returncode != 0:
                    subprocess.run(["git", "init"], cwd=self.path, check=True, capture_output=True)
            return None

    def commit(self, message: str, actor: str, paths: list[str], allow_empty: bool = False) -> str:
        repo = self.initialize()
        if repo is not None:
            from git import Actor

            identity = Actor(actor, f"{actor}@docpilot.local")
            repo.index.add(paths)
            commit = repo.index.commit(message, author=identity, committer=identity, skip_hooks=True)
            return commit.hexsha
        self._run("add", "--", *paths, actor=actor)
        args = ["commit", "-m", message]
        if allow_empty:
            args.insert(1, "--allow-empty")
        self._run(*args, actor=actor)
        return self._run("rev-parse", "HEAD", actor=actor).stdout.strip()

    def tag(self, tag: str, actor: str) -> None:
        repo = self.initialize()
        if repo is not None:
            with repo.git.custom_environment(
                GIT_COMMITTER_NAME=actor,
                GIT_COMMITTER_EMAIL=f"{actor}@docpilot.local",
            ):
                repo.create_tag(tag, message=f"DocPilot baseline {tag}")
            return
        self._run("tag", "-a", tag, "-m", f"DocPilot baseline {tag}", actor=actor)


class DocPilotService:
    def __init__(self, database: Database, root: Path):
        self.database = database
        self.root = Path(root)
        self.template_path = self.root / "templates" / "conops-pump.template.json"

    def _template(self) -> dict[str, Any]:
        return json.loads(self.template_path.read_text(encoding="utf-8"))

    def _grammar(self, name: str) -> str:
        path = self.root / "templates" / name
        if not path.is_file():
            raise DocPilotError(500, "GRAMMAR_NOT_FOUND", "Грамматика StrictDoc не найдена", [str(path)])
        return path.read_text(encoding="utf-8")

    def create_project(self, name: str) -> dict[str, Any]:
        project_id = f"project-{uuid.uuid4().hex[:8]}"
        repo_path = (self.root / "backend" / "data" / "repos" / project_id).resolve()
        with self.database.connect() as connection:
            connection.execute(
                "INSERT INTO project(id, name, git_path) VALUES (?, ?, ?)",
                (project_id, name.strip() or "Новый проект", str(repo_path)),
            )
        GitRepository(repo_path).initialize()
        return self._project(project_id)

    def create_document(self, project_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._project(project_id)
        template = self._template()
        document_id = f"doc-{uuid.uuid4().hex[:8]}"
        code = payload.get("code") or "ConOps"
        if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_-]{0,63}", code):
            raise DocPilotError(400, "DOCUMENT_CODE_INVALID", "Код документа недопустим", [code])
        if any(character in str(payload.get("title") or "") for character in "\r\n"):
            raise DocPilotError(400, "DOCUMENT_TITLE_INVALID", "Заголовок документа должен быть однострочным")
        with self.database.connect() as connection:
            connection.execute(
                """INSERT INTO document
                   (id, project_id, template_code, code, title, owner, status)
                   VALUES (?, ?, ?, ?, ?, ?, 'Draft')""",
                (
                    document_id,
                    project_id,
                    payload.get("template_code", template["code"]),
                    code,
                    payload.get("title", template["title"]),
                    payload.get("owner", "si"),
                ),
            )
            for index, (no, section) in enumerate(template["sections"].items(), start=1):
                connection.execute(
                    "INSERT INTO section(id, document_id, no, title, sort_order) VALUES (?, ?, ?, ?, ?)",
                    (f"{document_id}-sec-{index}", document_id, no, section["title"], index),
                )
        self.export_sdoc(document_id)
        return self.get_document(document_id)

    def _project(self, project_id: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            row = connection.execute("SELECT * FROM project WHERE id = ?", (project_id,)).fetchone()
        if not row:
            raise DocPilotError(404, "PROJECT_NOT_FOUND", "Проект не найден", [project_id])
        return dict(row)

    def _document_row(self, document_id: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            row = connection.execute("SELECT * FROM document WHERE id = ?", (document_id,)).fetchone()
        if not row:
            raise DocPilotError(404, "DOCUMENT_NOT_FOUND", "Документ не найден", [document_id])
        return dict(row)

    def list_facts(self) -> list[dict[str, Any]]:
        with self.database.connect() as connection:
            rows = connection.execute("SELECT * FROM fact ORDER BY id").fetchall()
        return [dict(row) for row in rows]

    def list_entities(self, kind: str | None = None) -> list[dict[str, Any]]:
        with self.database.connect() as connection:
            if kind:
                rows = connection.execute("SELECT * FROM entity WHERE kind = ? ORDER BY code", (kind,)).fetchall()
            else:
                rows = connection.execute("SELECT * FROM entity ORDER BY kind, code").fetchall()
        result = []
        for row in rows:
            item = dict(row)
            item["fields"] = _json(item["fields"], {})
            result.append(item)
        return result

    def _documents_using_reference(self, reference_id: str) -> list[str]:
        with self.database.connect() as connection:
            rows = connection.execute(
                """SELECT e.ref_id, e.supports, s.document_id
                   FROM element e JOIN section s ON s.id=e.section_id"""
            ).fetchall()
        return sorted(
            {
                row["document_id"]
                for row in rows
                if row["ref_id"] == reference_id or reference_id in _json(row["supports"], [])
            }
        )

    def create_fact(self, payload: dict[str, Any]) -> dict[str, Any]:
        reasons = []
        if not payload.get("source_doc") or not payload.get("source_anchor"):
            reasons.append("Для факта обязательны документ-источник и якорь.")
        if payload.get("value_num") is not None and not payload.get("value_unit"):
            reasons.append("Числовой факт должен иметь единицу измерения.")
        if reasons:
            raise DocPilotError(400, "FACT_INVALID", "Факт не прошёл проверку", reasons)
        fact_id = payload.get("id") or f"fact-{uuid.uuid4().hex[:8]}"
        try:
            with self.database.connect() as connection:
                connection.execute(
                    """INSERT INTO fact
                       (id, subject, predicate, value_num, value_unit, value_text, source_doc,
                        source_anchor, mark, disposition, author, at, version)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)""",
                    (
                        fact_id,
                        payload["subject"],
                        payload["predicate"],
                        payload.get("value_num"),
                        payload.get("value_unit"),
                        payload.get("value_text"),
                        payload["source_doc"],
                        payload["source_anchor"],
                        payload.get("mark", "И"),
                        payload.get("disposition", "noted"),
                        payload.get("author", "eng"),
                        utcnow(),
                    ),
                )
        except sqlite3.IntegrityError as error:
            raise DocPilotError(400, "FACT_INVALID", "Факт не создан", [str(error)]) from error
        return next(fact for fact in self.list_facts() if fact["id"] == fact_id)

    def update_fact(self, fact_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        affected_documents = self._documents_using_reference(fact_id)
        for document_id in affected_documents:
            self._assert_output_clean(document_id)
        with self.database.connect() as connection:
            current = connection.execute("SELECT * FROM fact WHERE id = ?", (fact_id,)).fetchone()
            if not current:
                raise DocPilotError(404, "FACT_NOT_FOUND", "Факт не найден", [fact_id])
            merged = dict(current) | payload
            reasons = []
            if not merged.get("source_doc") or not merged.get("source_anchor"):
                reasons.append("Для факта обязательны документ-источник и якорь.")
            if merged.get("value_num") is not None and not merged.get("value_unit"):
                reasons.append("Числовой факт должен иметь единицу измерения.")
            if reasons:
                raise DocPilotError(400, "FACT_INVALID", "Факт не прошёл проверку", reasons)
            connection.execute(
                """UPDATE fact SET subject=?, predicate=?, value_num=?, value_unit=?, value_text=?,
                   source_doc=?, source_anchor=?, mark=?, disposition=?, author=?, at=?, version=version+1
                   WHERE id=?""",
                (
                    merged["subject"], merged["predicate"], merged.get("value_num"), merged.get("value_unit"),
                    merged.get("value_text"), merged["source_doc"], merged["source_anchor"], merged["mark"],
                    merged["disposition"], payload.get("author", merged["author"]), utcnow(), fact_id,
                ),
            )
        for document_id in affected_documents:
            self.export_sdoc(document_id)
        return next(fact for fact in self.list_facts() if fact["id"] == fact_id)

    def create_entity(self, payload: dict[str, Any]) -> dict[str, Any]:
        if not str(payload.get("code") or "").strip() or not str(payload.get("title") or "").strip():
            raise DocPilotError(400, "ENTITY_INVALID", "Сущность не создана", ["Код и название обязательны."])
        entity_id = payload.get("id") or f"ent-{uuid.uuid4().hex[:8]}"
        try:
            with self.database.connect() as connection:
                connection.execute(
                    "INSERT INTO entity(id, kind, code, title, fields, version) VALUES (?, ?, ?, ?, ?, 1)",
                    (entity_id, payload["kind"], payload["code"], payload["title"], json.dumps(payload.get("fields", {}), ensure_ascii=False)),
                )
        except sqlite3.IntegrityError as error:
            raise DocPilotError(400, "ENTITY_INVALID", "Сущность не создана", [str(error)]) from error
        return next(entity for entity in self.list_entities() if entity["id"] == entity_id)

    def update_entity(self, entity_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        affected_documents = self._documents_using_reference(entity_id)
        for document_id in affected_documents:
            self._assert_output_clean(document_id)
        with self.database.connect() as connection:
            row = connection.execute("SELECT * FROM entity WHERE id = ?", (entity_id,)).fetchone()
            if not row:
                raise DocPilotError(404, "ENTITY_NOT_FOUND", "Сущность не найдена", [entity_id])
            current = dict(row)
            current_fields = _json(current["fields"], {})
            next_fields = payload.get("fields", current_fields)
            if not str(payload.get("code", current["code"])).strip() or not str(payload.get("title", current["title"])).strip():
                raise DocPilotError(400, "ENTITY_INVALID", "Сущность не изменена", ["Код и название обязательны."])
            version_delta = int(next_fields != current_fields)
            try:
                connection.execute(
                    """UPDATE entity SET kind=?, code=?, title=?, fields=?, version=version+? WHERE id=?""",
                    (
                        payload.get("kind", current["kind"]), payload.get("code", current["code"]),
                        payload.get("title", current["title"]), json.dumps(next_fields, ensure_ascii=False),
                        version_delta, entity_id,
                    ),
                )
            except sqlite3.IntegrityError as error:
                raise DocPilotError(400, "ENTITY_INVALID", "Сущность не изменена", [str(error)]) from error
        for document_id in affected_documents:
            self.export_sdoc(document_id)
        return next(entity for entity in self.list_entities() if entity["id"] == entity_id)

    def delete_fact(self, fact_id: str) -> None:
        self._delete_reference("fact", fact_id)

    def delete_entity(self, entity_id: str) -> None:
        self._delete_reference("entity", entity_id)

    def _delete_reference(self, table: str, reference_id: str) -> None:
        if table not in {"fact", "entity"}:
            raise ValueError(table)
        with self.database.connect() as connection:
            used = connection.execute("SELECT id FROM element WHERE ref_id = ? LIMIT 1", (reference_id,)).fetchone()
            support_rows = connection.execute("SELECT id, supports FROM element WHERE supports != '[]'").fetchall()
            support = next((row["id"] for row in support_rows if reference_id in _json(row["supports"], [])), None)
            if used or support:
                raise DocPilotError(
                    409,
                    "REFERENCE_IN_USE",
                    "Запись используется в документе",
                    [f"Элемент: {(used or {'id': support})['id']}"],
                )
            cursor = connection.execute(f"DELETE FROM {table} WHERE id = ?", (reference_id,))
            if cursor.rowcount == 0:
                code = "FACT_NOT_FOUND" if table == "fact" else "ENTITY_NOT_FOUND"
                raise DocPilotError(404, code, "Запись не найдена", [reference_id])

    def _resolve_element(self, connection: sqlite3.Connection, row: sqlite3.Row) -> dict[str, Any]:
        item = dict(row)
        item["supports"] = _json(item.pop("supports"), [])
        item["query"] = _json(item.pop("query_json"), None)
        item["resolved"] = None
        item["resolved_supports"] = []
        if item["kind"] == "fact_ref" and item["ref_id"]:
            resolved = connection.execute("SELECT * FROM fact WHERE id = ?", (item["ref_id"],)).fetchone()
            item["resolved"] = _dict(resolved)
        elif item["kind"] == "entity_ref" and item["ref_id"]:
            resolved = connection.execute("SELECT * FROM entity WHERE id = ?", (item["ref_id"],)).fetchone()
            if resolved:
                item["resolved"] = dict(resolved)
                item["resolved"]["fields"] = _json(item["resolved"]["fields"], {})
        for support_id in item["supports"]:
            support = connection.execute("SELECT * FROM fact WHERE id = ?", (support_id,)).fetchone()
            support_type = "fact"
            if support is None:
                support = connection.execute("SELECT * FROM entity WHERE id = ?", (support_id,)).fetchone()
                support_type = "entity"
            if support:
                resolved_support = dict(support)
                if support_type == "entity":
                    resolved_support["fields"] = _json(resolved_support["fields"], {})
                resolved_support["_type"] = support_type
                item["resolved_supports"].append(resolved_support)
        return item

    def _snapshot(self, document_id: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            document = connection.execute(
                """SELECT d.*, p.name AS project_name, p.git_path
                   FROM document d JOIN project p ON p.id = d.project_id WHERE d.id = ?""",
                (document_id,),
            ).fetchone()
            if not document:
                raise DocPilotError(404, "DOCUMENT_NOT_FOUND", "Документ не найден", [document_id])
            result = dict(document)
            sections = connection.execute(
                "SELECT * FROM section WHERE document_id = ? ORDER BY sort_order", (document_id,)
            ).fetchall()
            result["sections"] = []
            for section_row in sections:
                section = dict(section_row)
                elements = connection.execute(
                    "SELECT * FROM element WHERE section_id = ? ORDER BY sort_order", (section["id"],)
                ).fetchall()
                section["elements"] = [self._resolve_element(connection, row) for row in elements]
                result["sections"].append(section)
        return result

    @staticmethod
    def _matches_expectation(element: dict[str, Any], expectation: dict[str, Any]) -> bool:
        if element["kind"] != expectation["kind"]:
            return False
        if expectation.get("entity_kind"):
            return (element.get("resolved") or {}).get("kind") == expectation["entity_kind"]
        return True

    def get_document(self, document_id: str, stage: str = "SRR") -> dict[str, Any]:
        stage = stage.upper()
        if stage not in {"MCR", "SRR"}:
            raise DocPilotError(400, "STAGE_INVALID", "Неизвестная ступень полноты", [stage])
        document = self._snapshot(document_id)
        template = self._template()
        complete_sections = 0
        for section in document["sections"]:
            expectations = template["sections"][section["no"]]["expects"][stage]
            checks = []
            for expectation in expectations:
                count = sum(self._matches_expectation(element, expectation) for element in section["elements"])
                minimum = expectation["min"]
                missing = max(0, minimum - count)
                label = expectation.get("label") or expectation["kind"]
                checks.append(
                    {
                        "kind": expectation["kind"],
                        "entity_kind": expectation.get("entity_kind"),
                        "label": label,
                        "count": count,
                        "min": minimum,
                        "missing": missing,
                        "satisfied": missing == 0,
                        "message": None if missing == 0 else f"Нужно: ещё {missing} {label} → добавить",
                    }
                )
            achieved = sum(min(item["count"], item["min"]) for item in checks)
            required = sum(item["min"] for item in checks)
            complete = all(item["satisfied"] for item in checks)
            complete_sections += int(complete)
            section["completeness"] = {
                "stage": stage,
                "achieved": achieved,
                "required": required,
                "complete": complete,
                "checks": checks,
            }

        document["stage"] = stage
        document["completeness"] = {
            "complete_sections": complete_sections,
            "total_sections": len(document["sections"]),
            "complete": complete_sections == len(document["sections"]),
        }
        document["support_drift"] = self.support_drift(document_id)
        document["style"] = template["style"]
        return document

    def support_drift(self, document_id: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            baseline_row = connection.execute(
                "SELECT snapshot_json FROM baseline WHERE document_id = ? ORDER BY rowid DESC LIMIT 1",
                (document_id,),
            ).fetchone()
            if not baseline_row:
                return {"changed": False, "items": []}
            baseline_snapshot = _json(baseline_row["snapshot_json"], {})
            live_elements = {
                row["id"]: row["no"]
                for row in connection.execute(
                    """SELECT e.id, s.no FROM element e JOIN section s ON s.id=e.section_id
                       WHERE s.document_id=?""",
                    (document_id,),
                ).fetchall()
            }
            items = []
            seen: set[tuple[str, str]] = set()
            for section in baseline_snapshot.get("sections", []):
                for element in section.get("elements", []):
                    references: list[tuple[str, dict[str, Any]]] = []
                    if element.get("resolved"):
                        ref_type = "fact" if element.get("kind") == "fact_ref" else "entity"
                        references.append((ref_type, element["resolved"]))
                    references.extend(
                        (support.get("_type", "fact"), support)
                        for support in element.get("resolved_supports", [])
                    )
                    for ref_type, reference in references:
                        key = (element["id"], reference["id"])
                        if key in seen or element["id"] not in live_elements:
                            continue
                        seen.add(key)
                        row = connection.execute(
                            f"SELECT version FROM {ref_type} WHERE id = ?", (reference["id"],)
                        ).fetchone()
                        if row and row["version"] != reference.get("version"):
                            items.append(
                                {
                                    "element_id": element["id"],
                                    "ref_id": reference["id"],
                                    "section": live_elements[element["id"]],
                                    "baseline_version": reference.get("version"),
                                    "current_version": row["version"],
                                }
                            )
        return {"changed": bool(items), "items": items}

    def _assert_valid_element(self, payload: dict[str, Any]) -> None:
        kind = payload.get("kind")
        allowed = {"fact_ref", "entity_ref", "statement", "query", "figure", "table"}
        if kind not in allowed:
            raise DocPilotError(400, "ELEMENT_KIND_INVALID", "Неизвестный вид элемента", [str(kind)])
        reasons: list[str] = []
        author = str(payload.get("author") or "eng")
        if author not in {"si", "eng", "rev"}:
            reasons.append("Автор должен быть одной из ролей: si, eng, rev.")
        if kind in {"fact_ref", "entity_ref"} and not payload.get("ref_id"):
            reasons.append("Ссылка выбирается из пикера и обязательна.")
        if kind in {"fact_ref", "entity_ref", "figure"} and payload.get("ref_id") and not re.fullmatch(
            r"[A-Za-z0-9][A-Za-z0-9_:/.-]{0,255}", str(payload["ref_id"])
        ):
            reasons.append("Ссылка REF содержит недопустимые символы.")
        if kind == "statement":
            text = (payload.get("text") or "").strip()
            if not text:
                reasons.append("Текст тезиса обязателен.")
            sentence_count = len([part for part in re.split(r"[.!?]+", text) if part.strip()])
            if sentence_count > 3:
                reasons.append("Тезис должен состоять из 1–3 предложений.")
            if extract_numbers(text) and not payload.get("supports"):
                reasons.append("Число без основания: выберите хотя бы одну опору.")
        if kind in {"query", "table"}:
            query = payload.get("query")
            if (
                not isinstance(query, dict)
                or query.get("kind") not in {"stakeholder", "scenario", "mode", "requirement"}
                or not str(query.get("filter") or "").strip()
                or not isinstance(query.get("columns"), list)
                or not query["columns"]
                or any(not isinstance(column, str) or not column.strip() for column in query["columns"])
            ):
                reasons.append("Для запроса обязательны вид, фильтр и непустой список колонок.")
            elif any("\n" in value or "\r" in value for value in [str(query["kind"]), str(query["filter"]), *query["columns"]]):
                reasons.append("Поля запроса должны быть однострочными.")
        if kind == "figure":
            if not str(payload.get("ref_id") or "").strip():
                reasons.append("Для фигуры обязательна ссылка REF.")
            if not str(payload.get("text") or "").strip():
                reasons.append("Для фигуры обязательна подпись CAPTION.")
            elif any(character in str(payload["text"]) for character in "\r\n"):
                reasons.append("Подпись фигуры должна быть однострочной.")
        if reasons:
            raise DocPilotError(400, "ELEMENT_INVALID", "Элемент не прошёл проверку", reasons)

    def add_element(self, document_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._assert_output_clean(document_id)
        self._assert_valid_element(payload)
        with self.database.connect() as connection:
            section = connection.execute(
                "SELECT * FROM section WHERE id = ? AND document_id = ?", (payload["section_id"], document_id)
            ).fetchone()
            if not section:
                raise DocPilotError(404, "SECTION_NOT_FOUND", "Раздел не найден", [payload["section_id"]])
            ref_version = None
            if payload["kind"] == "fact_ref":
                ref = connection.execute("SELECT version FROM fact WHERE id = ?", (payload.get("ref_id"),)).fetchone()
                if not ref:
                    raise DocPilotError(400, "REFERENCE_INVALID", "Факт должен быть выбран из базы", [str(payload.get("ref_id"))])
                ref_version = ref["version"]
            elif payload["kind"] == "entity_ref":
                ref = connection.execute("SELECT version FROM entity WHERE id = ?", (payload.get("ref_id"),)).fetchone()
                if not ref:
                    raise DocPilotError(400, "REFERENCE_INVALID", "Сущность должна быть выбрана из базы", [str(payload.get("ref_id"))])
                ref_version = ref["version"]
            for support in payload.get("supports", []):
                exists = connection.execute("SELECT 1 FROM fact WHERE id = ?", (support,)).fetchone() or connection.execute(
                    "SELECT 1 FROM entity WHERE id = ?", (support,)
                ).fetchone()
                if not exists:
                    raise DocPilotError(400, "SUPPORT_INVALID", "Опора должна быть выбрана из базы", [support])
            maximum = connection.execute("SELECT COALESCE(MAX(sort_order), 0) FROM element WHERE section_id = ?", (section["id"],)).fetchone()[0]
            numeric_ids = [
                int(match.group(1))
                for row in connection.execute("SELECT id FROM element").fetchall()
                if (match := re.fullmatch(r"EL-(\d+)", row["id"]))
            ]
            element_id = f"EL-{max(numeric_ids, default=0) + 1:03d}"
            connection.execute(
                """INSERT INTO element
                   (id, section_id, kind, ref_id, ref_version, text, supports, query_json,
                    author, at, sort_order, version)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)""",
                (
                    element_id, section["id"], payload["kind"], payload.get("ref_id"), ref_version,
                    payload.get("text"), json.dumps(payload.get("supports", [])),
                    json.dumps(payload.get("query"), ensure_ascii=False) if payload.get("query") else None,
                    payload.get("author", "eng"), utcnow(), maximum + 1,
                ),
            )
        self.export_sdoc(document_id)
        return next(
            element
            for section in self._snapshot(document_id)["sections"]
            for element in section["elements"]
            if element["id"] == element_id
        )

    def update_element(self, document_id: str, element_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        self._assert_output_clean(document_id)
        with self.database.connect() as connection:
            row = connection.execute(
                """SELECT e.* FROM element e JOIN section s ON s.id=e.section_id
                   WHERE e.id=? AND s.document_id=?""", (element_id, document_id)
            ).fetchone()
            if not row:
                raise DocPilotError(404, "ELEMENT_NOT_FOUND", "Элемент не найден", [element_id])
            current = dict(row)
            merged = {
                "kind": current["kind"],
                "ref_id": current["ref_id"],
                "text": current["text"],
                "supports": _json(current["supports"], []),
                "query": _json(current["query_json"], None),
                "author": current["author"],
            } | payload
            self._assert_valid_element(merged)
            for support in merged.get("supports", []):
                exists = connection.execute("SELECT 1 FROM fact WHERE id = ?", (support,)).fetchone() or connection.execute(
                    "SELECT 1 FROM entity WHERE id = ?", (support,)
                ).fetchone()
                if not exists:
                    raise DocPilotError(400, "SUPPORT_INVALID", "Опора должна быть выбрана из базы", [support])
            connection.execute(
                """UPDATE element SET text=?, supports=?, query_json=?, author=?, at=?, version=version+1 WHERE id=?""",
                (
                    merged.get("text"), json.dumps(merged.get("supports", [])),
                    json.dumps(merged.get("query"), ensure_ascii=False) if merged.get("query") else None,
                    payload.get("author", current["author"]), utcnow(), element_id,
                ),
            )
        self.export_sdoc(document_id)
        return next(
            element
            for section in self._snapshot(document_id)["sections"]
            for element in section["elements"]
            if element["id"] == element_id
        )

    def delete_element(self, document_id: str, element_id: str) -> None:
        self._assert_output_clean(document_id)
        with self.database.connect() as connection:
            cursor = connection.execute(
                """DELETE FROM element WHERE id = ? AND section_id IN
                   (SELECT id FROM section WHERE document_id = ?)""",
                (element_id, document_id),
            )
            if cursor.rowcount == 0:
                raise DocPilotError(404, "ELEMENT_NOT_FOUND", "Элемент не найден", [element_id])
        self.export_sdoc(document_id)

    def _sdoc_path(self, document: dict[str, Any]) -> Path:
        code = str(document["code"])
        if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_-]{0,63}", code):
            raise DocPilotError(500, "DOCUMENT_CODE_INVALID", "Код документа недопустим", [code])
        docs_root = (Path(document["git_path"]) / "docs").resolve()
        path = (docs_root / f"{code}.sdoc").resolve()
        if docs_root not in path.parents:
            raise DocPilotError(500, "DOCUMENT_PATH_INVALID", "Путь документа выходит за границы репозитория")
        return path

    def _assert_output_clean(self, document_id: str) -> None:
        document = self._document_row(document_id)
        stored_hash = document.get("generated_sdoc_hash")
        if not stored_hash:
            return
        project = self._project(document["project_id"])
        path = Path(project["git_path"]) / "docs" / f"{document['code']}.sdoc"
        if not path.exists() or fingerprint(path.read_text(encoding="utf-8")) != stored_hash:
            raise DocPilotError(
                409,
                "SDOC_DIVERGED",
                "Файл расходится с базой",
                ["docs/ConOps.sdoc изменён или удалён вне DocPilot; восстановите файл из Git."],
            )

    def _assert_rendered_output_clean(self, document_id: str) -> None:
        document = self._document_row(document_id)
        stored_hash = document.get("generated_rendered_hash")
        if not stored_hash:
            return
        project = self._project(document["project_id"])
        path = Path(project["git_path"]) / "docs" / "rendered" / f"{document['code']}.sdoc"
        if not path.exists() or fingerprint(path.read_text(encoding="utf-8")) != stored_hash:
            raise DocPilotError(
                409,
                "RENDERED_SDOC_DIVERGED",
                "Файл изложенного документа расходится с базой",
                ["docs/rendered/ConOps.sdoc изменён или удалён вне DocPilot; восстановите файл из Git."],
            )

    def _remember_rendered_output(self, document_id: str, content: str) -> None:
        with self.database.connect() as connection:
            connection.execute(
                "UPDATE document SET generated_rendered_hash = ? WHERE id = ?",
                (fingerprint(content), document_id),
            )

    def export_sdoc(self, document_id: str) -> dict[str, Any]:
        self._assert_output_clean(document_id)
        document = self._snapshot(document_id)
        content = generate_structure_sdoc(document, self._grammar("orbita.sgra"))
        path = self._sdoc_path(document)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        typst_path = Path(document["git_path"]) / "docs" / "typst" / "conops.typ"
        typst_path.parent.mkdir(parents=True, exist_ok=True)
        typst_path.write_text((self.root / "templates" / "conops.typ").read_text(encoding="utf-8"), encoding="utf-8")
        digest = fingerprint(content)
        with self.database.connect() as connection:
            connection.execute("UPDATE document SET generated_sdoc_hash = ? WHERE id = ?", (digest, document_id))
        return {"path": str(path), "typst_path": str(typst_path), "content": content, "sha256": digest}

    def _verify_sdoc_path(self, path: Path) -> dict[str, Any]:
        executable = find_tool("strictdoc")
        if not executable:
            return {
                "valid": None,
                "verified": False,
                "status": "skipped",
                "mode": "internal-structure",
                "reason": "STRICTDOC_NOT_INSTALLED",
                "message": "Файл сгенерирован, но совместимость не подтверждена: CLI StrictDoc не установлен.",
            }
        help_result = subprocess.run([executable, "-h"], text=True, capture_output=True, check=False)
        supports_check = "check" in help_result.stdout.split("command:", 1)[-1].split("Further help:", 1)[0]
        if supports_check:
            command = [executable, "check", str(path)]
            mode = "strictdoc-check"
            temporary = None
        else:
            temporary = tempfile.TemporaryDirectory(prefix="docpilot-strictdoc-")
            command = [
                executable,
                "export",
                str(path),
                "--formats",
                "json",
                "--output-dir",
                temporary.name,
                "--no-parallelization",
            ]
            mode = "strictdoc-export-parse"
        result = subprocess.run(command, text=True, capture_output=True, check=False)
        if temporary is not None:
            temporary.cleanup()
        if result.returncode != 0:
            raise DocPilotError(422, "STRICTDOC_INVALID", "StrictDoc отклонил документ", [result.stderr or result.stdout])
        return {
            "valid": True,
            "verified": True,
            "status": "passed",
            "mode": mode,
            "message": "StrictDoc разобрал и экспортировал документ без ошибок." if mode == "strictdoc-export-parse" else "strictdoc check: OK",
        }

    def _verify_rendered_content(self, content: str) -> None:
        with tempfile.TemporaryDirectory(prefix="docpilot-rendered-check-") as temporary:
            path = Path(temporary) / "rendered.sdoc"
            path.write_text(content, encoding="utf-8")
            self._verify_sdoc_path(path)

    def check_sdoc(self, document_id: str) -> dict[str, Any]:
        exported = self.export_sdoc(document_id)
        result = self._verify_sdoc_path(Path(exported["path"]))
        result["sha256"] = exported["sha256"]
        return result

    @staticmethod
    def _baseline_items(document: dict[str, Any]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        seen_refs: set[tuple[str, str]] = set()
        for section in document["sections"]:
            for element in section["elements"]:
                items.append({"type": "element", "id": element["id"], "version": element["version"]})
                resolved = element.get("resolved")
                ref_type = "fact" if element["kind"] == "fact_ref" else "entity"
                if resolved and (ref_type, resolved["id"]) not in seen_refs:
                    items.append({"type": ref_type, "id": resolved["id"], "version": resolved["version"]})
                    seen_refs.add((ref_type, resolved["id"]))
                for support in element.get("resolved_supports", []):
                    ref_type = support.get("_type", "fact")
                    key = (ref_type, support["id"])
                    if key not in seen_refs:
                        items.append({"type": ref_type, "id": support["id"], "version": support["version"]})
                        seen_refs.add(key)
        return items

    def create_baseline(self, document_id: str, name: str, actor: str = "si") -> dict[str, Any]:
        if actor != "si":
            raise DocPilotError(403, "ROLE_FORBIDDEN", "Базировать документ может только ведущий СИ", [actor])
        clean_name = re.sub(r"[^A-Za-z0-9_-]+", "-", name.strip().upper()).strip("-") or "BASELINE"
        verification = self.check_sdoc(document_id)
        if not verification.get("verified"):
            raise DocPilotError(
                503,
                "STRICTDOC_UNAVAILABLE",
                "Базирование требует проверку StrictDoc",
                [verification.get("message", "StrictDoc CLI недоступен")],
            )
        document = self._snapshot(document_id)
        with self.database.connect() as connection:
            count = connection.execute("SELECT COUNT(*) FROM baseline WHERE document_id = ? AND name = ?", (document_id, clean_name)).fetchone()[0]
        tag = f"BL-{clean_name}-{count + 1}"
        repo = GitRepository(Path(document["git_path"]))
        relative = str(self._sdoc_path(document).relative_to(document["git_path"]))
        typst_relative = str((Path(document["git_path"]) / "docs" / "typst" / "conops.typ").relative_to(document["git_path"]))
        commit_hash = repo.commit(f"baseline: {tag}", actor, [relative, typst_relative], allow_empty=True)
        repo.tag(tag, actor)
        baseline_id = f"baseline-{uuid.uuid4().hex[:10]}"
        snapshot = document
        items = self._baseline_items(snapshot)
        at = utcnow()
        with self.database.connect() as connection:
            connection.execute(
                """INSERT INTO baseline
                   (id, document_id, name, git_tag, items, snapshot_json, by_user, at, commit_hash)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    baseline_id, document_id, clean_name, tag, json.dumps(items, ensure_ascii=False),
                    json.dumps(snapshot, ensure_ascii=False), actor, at, commit_hash,
                ),
            )
            connection.execute("UPDATE document SET status = 'Baseline' WHERE id = ?", (document_id,))
        logger.info("baseline created tag=%s commit=%s actor=%s", tag, commit_hash, actor)
        return {
            "id": baseline_id,
            "document_id": document_id,
            "name": clean_name,
            "git_tag": tag,
            "items": items,
            "by": actor,
            "at": at,
            "commit_hash": commit_hash,
            "authors": self._authors(snapshot, None, ""),
        }

    def list_baselines(self, document_id: str) -> list[dict[str, Any]]:
        with self.database.connect() as connection:
            rows = connection.execute("SELECT * FROM baseline WHERE document_id = ? ORDER BY at DESC, rowid DESC", (document_id,)).fetchall()
        result = []
        for row in rows:
            snapshot = _json(row["snapshot_json"], {})
            result.append(
                {
                    "id": row["id"], "document_id": row["document_id"], "name": row["name"],
                    "git_tag": row["git_tag"], "items": _json(row["items"], []), "by": row["by_user"],
                    "at": row["at"], "commit_hash": row["commit_hash"],
                    "authors": self._authors(snapshot, None, ""),
                }
            )
        return result

    def _baseline(self, baseline_id_or_tag: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            row = connection.execute(
                "SELECT * FROM baseline WHERE id = ? OR git_tag = ?", (baseline_id_or_tag, baseline_id_or_tag)
            ).fetchone()
        if not row:
            raise DocPilotError(404, "BASELINE_NOT_FOUND", "Базирование не найдено", [baseline_id_or_tag])
        result = dict(row)
        result["items"] = _json(result["items"], [])
        result["snapshot"] = _json(result.pop("snapshot_json"), {})
        result["by"] = result.pop("by_user")
        return result

    @staticmethod
    def _element_map(snapshot: dict[str, Any]) -> dict[str, dict[str, Any]]:
        return {
            element["id"]: element
            for section in snapshot["sections"]
            for element in section["elements"]
        }

    def diff_baselines(self, from_id: str, to_id: str) -> dict[str, Any]:
        before = self._baseline(from_id)
        after = self._baseline(to_id)
        if before["document_id"] != after["document_id"]:
            raise DocPilotError(400, "DIFF_SCOPE_INVALID", "Базирования относятся к разным документам")
        left = self._element_map(before["snapshot"])
        right = self._element_map(after["snapshot"])
        added = [
            {"mid": mid, "element": right[mid], "author": after["by"], "element_author": right[mid]["author"]}
            for mid in sorted(right.keys() - left.keys())
        ]
        removed = [
            {"mid": mid, "element": left[mid], "author": before["by"], "element_author": left[mid]["author"]}
            for mid in sorted(left.keys() - right.keys())
        ]
        ignored = {"at", "resolved", "resolved_supports", "version"}
        changed = []
        for mid in sorted(left.keys() & right.keys()):
            fields = []
            keys = (set(left[mid]) | set(right[mid])) - ignored
            for key in sorted(keys):
                if left[mid].get(key) != right[mid].get(key):
                    label = "STATEMENT" if key == "text" else key.upper()
                    fields.append({"field": label, "from": left[mid].get(key), "to": right[mid].get(key)})
            before_refs = []
            after_refs = []
            if left[mid].get("resolved"):
                before_refs.append(left[mid]["resolved"])
            if right[mid].get("resolved"):
                after_refs.append(right[mid]["resolved"])
            before_refs.extend(left[mid].get("resolved_supports", []))
            after_refs.extend(right[mid].get("resolved_supports", []))
            before_ref_map = {reference["id"]: reference for reference in before_refs}
            after_ref_map = {reference["id"]: reference for reference in after_refs}
            for ref_id in sorted(set(before_ref_map) & set(after_ref_map)):
                old_ref = before_ref_map[ref_id]
                new_ref = after_ref_map[ref_id]
                for key in sorted((set(old_ref) | set(new_ref)) - {"id", "version", "at", "_type"}):
                    if old_ref.get(key) != new_ref.get(key):
                        fields.append(
                            {
                                "field": f"REF.{ref_id}.{key.upper()}",
                                "from": old_ref.get(key),
                                "to": new_ref.get(key),
                            }
                        )
            if fields:
                changed.append(
                    {"mid": mid, "fields": fields, "author": after["by"], "element_author": right[mid]["author"]}
                )
        return {
            "from": before["git_tag"], "to": after["git_tag"],
            "summary": {"added": len(added), "changed": len(changed), "removed": len(removed)},
            "added": added, "changed": changed, "removed": removed,
            "strategy": "application-level MID diff",
        }

    @staticmethod
    def _allowed_numbers(element: dict[str, Any]) -> set[str]:
        values: list[str] = []
        if element.get("text"):
            values.append(element["text"])
        if element.get("query"):
            values.append(json.dumps(element["query"], ensure_ascii=False))
        resolved = element.get("resolved") or {}
        for field in ("subject", "predicate", "value_text", "title"):
            if resolved.get(field) is not None:
                values.append(str(resolved[field]))
        if resolved.get("value_num") is not None:
            values.append(str(resolved["value_num"]))
            if resolved.get("value_unit"):
                values.append(f"{resolved['value_num']} {resolved['value_unit']}")
                value_text = str(resolved.get("value_text") or "").lower()
                if "не более" in value_text:
                    values.append(f"≤{resolved['value_num']} {resolved['value_unit']}")
                if "не менее" in value_text:
                    values.append(f"≥{resolved['value_num']} {resolved['value_unit']}")
        if resolved.get("fields"):
            values.append(json.dumps(resolved["fields"], ensure_ascii=False))
        return set(number for value in values for number in extract_numbers(value))

    def validate_rendered_sections(self, sections: list[dict[str, Any]], source: dict[str, Any]) -> None:
        source_sections = {section["no"]: section for section in source["sections"]}
        violations: list[str] = []
        for rendered in sections:
            source_section = source_sections.get(rendered["no"])
            if not source_section:
                violations.append(f"{rendered['no']}: раздел отсутствует в исходной структуре")
                continue
            elements = {element["id"]: element for element in source_section["elements"]}
            for line in rendered["text"].splitlines():
                if not line.strip():
                    continue
                markers = MARKER_RE.findall(line)
                invalid = [marker for marker in markers if marker not in elements]
                if invalid:
                    violations.append(f"{rendered['no']}: ссылки вне раздела: {', '.join(invalid)}")
                numbers = extract_numbers(line)
                if not markers:
                    violations.append(f"{rendered['no']}: утверждение без ссылки [MID]")
                if not numbers:
                    continue
                if not markers:
                    violations.extend(f"{rendered['no']}: число вне элементов: {number}" for number in numbers)
                    continue
                allowed = set().union(*(self._allowed_numbers(elements[mid]) for mid in markers if mid in elements))
                violations.extend(
                    f"{rendered['no']}: число вне элементов: {number}" for number in numbers if number not in allowed
                )
        if violations:
            logger.warning("render rejected code=RENDER_TRACEABILITY_FAILED reasons=%s", violations)
            raise DocPilotError(422, "RENDER_TRACEABILITY_FAILED", "Рендеринг содержит неподтверждённые данные", violations)

    @staticmethod
    def _stub_sentence(element: dict[str, Any]) -> str:
        mid = element["id"]
        if element["kind"] == "statement":
            return f"{element.get('text', '').strip()} [{mid}]"
        if element["kind"] == "fact_ref":
            fact = element.get("resolved") or {}
            value = fact.get("value_text") or ""
            if fact.get("value_num") is not None:
                value = " ".join(part for part in [value, _display_number(fact["value_num"]), fact.get("value_unit")] if part)
            return f"Для {fact.get('subject', 'объекта').lower()} принято: {fact.get('predicate', 'значение')} {value}. [{mid}]"
        if element["kind"] == "entity_ref":
            entity = element.get("resolved") or {}
            fields = entity.get("fields") or {}
            detail = fields.get("description") or fields.get("interest") or fields.get("outcome") or "учтено в концепции"
            return f"{entity.get('title', 'Сущность')}: {detail}. [{mid}]"
        if element["kind"] == "query":
            return f"Перечень требований формируется по актуальному реестру и заданному фильтру. [{mid}]"
        if element["kind"] == "figure":
            return f"Иллюстрация «{element.get('text') or 'без названия'}» включена в раздел. [{mid}]"
        return f"Табличное представление сформировано по запросу раздела. [{mid}]"

    def _build_stub_sections(self, source: dict[str, Any]) -> list[dict[str, Any]]:
        return [
            {
                "no": section["no"],
                "title": section["title"],
                "text": "\n".join(self._stub_sentence(element) for element in section["elements"]),
                "element_links": [
                    {"mids": [element["id"]], "label": element["id"], "source": element}
                    for element in section["elements"]
                ],
            }
            for section in source["sections"]
        ]

    @staticmethod
    def _render_text_diff(previous: dict[str, Any] | None, current: dict[str, Any]) -> dict[str, Any]:
        if previous is None:
            return {"from_rendering": None, "summary": {"changed_sections": 0}, "sections": []}
        before = {section["no"]: section for section in previous["sections"]}
        after = {section["no"]: section for section in current["sections"]}
        sections = []
        for no in sorted(set(before) | set(after)):
            old_text = before.get(no, {}).get("text", "")
            new_text = after.get(no, {}).get("text", "")
            if old_text == new_text:
                continue
            patch = "\n".join(
                difflib.unified_diff(
                    old_text.splitlines(), new_text.splitlines(),
                    fromfile=f"{previous['id']}:{no}", tofile=f"{current['id']}:{no}", lineterm="",
                )
            )
            sections.append({"no": no, "patch": patch})
        return {
            "from_rendering": previous["id"],
            "summary": {"changed_sections": len(sections)},
            "sections": sections,
        }

    def create_rendering(
        self,
        document_id: str,
        baseline_id: str | None = None,
        actor: str = "si",
        simulate_sections: dict[str, str] | None = None,
    ) -> dict[str, Any]:
        self._assert_rendered_output_clean(document_id)
        if baseline_id:
            baseline = self._baseline(baseline_id)
            if baseline["document_id"] != document_id:
                raise DocPilotError(400, "BASELINE_SCOPE_INVALID", "Базирование относится к другому документу")
            source = baseline["snapshot"]
            baseline_db_id = baseline["id"]
        else:
            source = self._snapshot(document_id)
            baseline_db_id = None
        template = self._template()
        prompt_payload = {"style": template["style"], "sections": source["sections"]}
        prompt_hash = fingerprint(json.dumps(prompt_payload, ensure_ascii=False, sort_keys=True))[:16]
        sections = self._build_stub_sections(source)
        if simulate_sections:
            for section in sections:
                if section["no"] in simulate_sections:
                    section["text"] = simulate_sections[section["no"]]
        self.validate_rendered_sections(sections, source)
        rendering_id = f"render-{uuid.uuid4().hex[:10]}"
        with self.database.connect() as connection:
            version = connection.execute("SELECT COUNT(*) FROM rendering WHERE document_id = ?", (document_id,)).fetchone()[0] + 1
        rendering = {
            "id": rendering_id,
            "document_id": document_id,
            "baseline_id": baseline_db_id,
            "sections": sections,
            "engine": "stub",
            "model": None,
            "prompt_fingerprint": prompt_hash,
            "patches": [],
            "reviewer": None,
            "accepted_at": None,
            "created_at": utcnow(),
            "version": version,
        }
        with self.database.connect() as connection:
            previous_row = connection.execute(
                "SELECT id FROM rendering WHERE document_id = ? ORDER BY rowid DESC LIMIT 1",
                (document_id,),
            ).fetchone()
        previous = self._rendering(previous_row["id"]) if previous_row else None
        rendering["text_diff"] = self._render_text_diff(previous, rendering)
        rendered_content = generate_rendered_sdoc(
            source["title"], rendering, self._grammar("rendered.sgra")
        )
        self._verify_rendered_content(rendered_content)
        with self.database.connect() as connection:
            connection.execute(
                """INSERT INTO rendering
                   (id, document_id, baseline_id, sections, engine, model, prompt_fingerprint,
                    patches, reviewer, accepted_at, created_at, version)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    rendering_id, document_id, baseline_db_id, json.dumps(sections, ensure_ascii=False), "stub", None,
                    prompt_hash, "[]", None, None, rendering["created_at"], version,
                ),
            )
        document = self._snapshot(document_id)
        repo_path = Path(document["git_path"])
        rendered_path = repo_path / "docs" / "rendered" / f"{document['code']}.sdoc"
        rendered_path.parent.mkdir(parents=True, exist_ok=True)
        rendered_path.write_text(rendered_content, encoding="utf-8")
        self._remember_rendered_output(document_id, rendered_content)
        GitRepository(repo_path).commit(
            f"render: v{version} ({rendering['engine']}, {prompt_hash})",
            actor,
            [str(rendered_path.relative_to(repo_path))],
        )
        logger.info("rendering created id=%s engine=%s fingerprint=%s", rendering_id, rendering["engine"], prompt_hash)
        return rendering

    def _rendering(self, rendering_id: str) -> dict[str, Any]:
        with self.database.connect() as connection:
            row = connection.execute("SELECT * FROM rendering WHERE id = ?", (rendering_id,)).fetchone()
        if not row:
            raise DocPilotError(404, "RENDERING_NOT_FOUND", "Рендеринг не найден", [rendering_id])
        result = dict(row)
        result["sections"] = _json(result["sections"], [])
        result["patches"] = _json(result["patches"], [])
        return result

    def list_renderings(self, document_id: str) -> list[dict[str, Any]]:
        with self.database.connect() as connection:
            ids = [row["id"] for row in connection.execute(
                "SELECT id FROM rendering WHERE document_id = ? ORDER BY rowid", (document_id,)
            ).fetchall()]
        renderings = [self._rendering(id_) for id_ in ids]
        previous: dict[str, Any] | None = None
        for rendering in renderings:
            rendering["text_diff"] = self._render_text_diff(previous, rendering)
            previous = rendering
        return list(reversed(renderings))

    @staticmethod
    def _patch(old: str, new: str) -> str:
        try:
            from diff_match_patch import diff_match_patch

            dmp = diff_match_patch()
            return dmp.patch_toText(dmp.patch_make(old, new))
        except ImportError:
            return "\n".join(difflib.unified_diff(old.splitlines(), new.splitlines(), lineterm=""))

    def update_rendering_section(self, rendering_id: str, no: str, text: str, reviewer: str = "rev") -> dict[str, Any]:
        if reviewer != "rev":
            raise DocPilotError(403, "ROLE_FORBIDDEN", "Изложение может править только рецензент", [reviewer])
        rendering = self._rendering(rendering_id)
        self._assert_rendered_output_clean(rendering["document_id"])
        if rendering["accepted_at"]:
            raise DocPilotError(409, "RENDERING_ACCEPTED", "Принятый рендеринг нельзя изменять")
        target = next((section for section in rendering["sections"] if section["no"] == no), None)
        if not target:
            raise DocPilotError(404, "RENDER_SECTION_NOT_FOUND", "Раздел рендеринга не найден", [no])
        old_text = target["text"]
        reasons = []
        old_markers = set(MARKER_RE.findall(old_text))
        new_markers = set(MARKER_RE.findall(text))
        if old_markers != new_markers:
            missing = sorted(old_markers - new_markers)
            added = sorted(new_markers - old_markers)
            if missing:
                reasons.append(f"Удалены ссылки на элементы: {', '.join(missing)}")
            if added:
                reasons.append(f"Добавлены несогласованные ссылки: {', '.join(added)}")
        old_numbers = sorted(extract_numbers(old_text))
        new_numbers = sorted(extract_numbers(text))
        if old_numbers != new_numbers:
            reasons.append(f"Изменены числовые данные: {', '.join(old_numbers) or 'нет'} → {', '.join(new_numbers) or 'нет'}")
        old_qualifiers = extract_qualifiers(old_text)
        new_qualifiers = extract_qualifiers(text)
        if old_qualifiers != new_qualifiers:
            reasons.append(
                "Изменены смысловые квалификаторы: "
                f"{', '.join(old_qualifiers) or 'нет'} → {', '.join(new_qualifiers) or 'нет'}"
            )
        if reasons:
            logger.warning("review rejected rendering=%s section=%s reasons=%s", rendering_id, no, reasons)
            raise DocPilotError(422, "REVIEW_SEMANTICS_CHANGED", "Правка меняет смысл или трассируемость", reasons)
        source = self._baseline(rendering["baseline_id"])["snapshot"] if rendering["baseline_id"] else self._snapshot(rendering["document_id"])
        candidate_sections = [dict(section) for section in rendering["sections"]]
        next_target = next(section for section in candidate_sections if section["no"] == no)
        next_target["text"] = text
        self.validate_rendered_sections(candidate_sections, source)
        document = self._snapshot(rendering["document_id"])
        candidate_rendering = dict(rendering)
        candidate_rendering["sections"] = candidate_sections
        rendered_content = generate_rendered_sdoc(
            document["title"], candidate_rendering, self._grammar("rendered.sgra")
        )
        self._verify_rendered_content(rendered_content)
        patch = {"section": no, "author": reviewer, "at": utcnow(), "patch": self._patch(old_text, text)}
        patches = [*rendering["patches"], patch]
        with self.database.connect() as connection:
            connection.execute(
                "UPDATE rendering SET sections = ?, patches = ?, reviewer = ? WHERE id = ?",
                (json.dumps(candidate_sections, ensure_ascii=False), json.dumps(patches, ensure_ascii=False), reviewer, rendering_id),
            )
        updated = self._rendering(rendering_id)
        repo_path = Path(document["git_path"])
        rendered_path = repo_path / "docs" / "rendered" / f"{document['code']}.sdoc"
        rendered_path.write_text(rendered_content, encoding="utf-8")
        self._remember_rendered_output(rendering["document_id"], rendered_content)
        GitRepository(repo_path).commit(
            f"review: {rendering_id} {no} by {reviewer}", reviewer, [str(rendered_path.relative_to(repo_path))]
        )
        return updated

    def accept_rendering(self, rendering_id: str, reviewer: str = "rev") -> dict[str, Any]:
        if reviewer != "rev":
            raise DocPilotError(403, "ROLE_FORBIDDEN", "Рендеринг может принять только рецензент", [reviewer])
        rendering = self._rendering(rendering_id)
        accepted_at = rendering["accepted_at"] or utcnow()
        with self.database.connect() as connection:
            connection.execute(
                "UPDATE rendering SET reviewer = ?, accepted_at = ? WHERE id = ?", (reviewer, accepted_at, rendering_id)
            )
        return self._rendering(rendering_id)

    @staticmethod
    def _escape_pdf_text(text: str) -> str:
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    @classmethod
    def _minimal_pdf(cls, title: str, lines: list[str]) -> bytes:
        display = [title, *lines]
        commands = ["BT", "/F1 16 Tf", "72 770 Td"]
        for index, line in enumerate(display):
            if index:
                commands.extend(["0 -24 Td", "/F1 11 Tf"])
            ascii_line = line.encode("ascii", "replace").decode("ascii")
            commands.append(f"({cls._escape_pdf_text(ascii_line)}) Tj")
        commands.append("ET")
        stream = "\n".join(commands).encode("ascii")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
        pdf = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, obj in enumerate(objects, start=1):
            offsets.append(len(pdf))
            pdf.extend(f"{index} 0 obj\n".encode() + obj + b"\nendobj\n")
        xref = len(pdf)
        pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode())
        pdf.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            pdf.extend(f"{offset:010d} 00000 n \n".encode())
        pdf.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
        return bytes(pdf)

    @classmethod
    def _fallback_pdf(cls, title: str, lines: list[str], sections: list[dict[str, Any]]) -> bytes:
        try:
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

            font_path = next(
                (
                    path for path in [
                        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
                        Path("/Library/Fonts/Arial.ttf"),
                    ]
                    if path.is_file()
                ),
                None,
            )
            font_name = "Helvetica"
            if font_path:
                font_name = "DocPilotSans"
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "DocPilotTitle", parent=styles["Title"], fontName=font_name,
                fontSize=21, leading=27, textColor="#17303a", alignment=TA_CENTER, spaceAfter=12,
            )
            meta_style = ParagraphStyle(
                "DocPilotMeta", parent=styles["BodyText"], fontName=font_name,
                fontSize=10, leading=15, textColor="#52666e", alignment=TA_CENTER,
            )
            heading_style = ParagraphStyle(
                "DocPilotHeading", parent=styles["Heading2"], fontName=font_name,
                fontSize=14, leading=18, textColor="#176b87", spaceBefore=12, spaceAfter=7,
            )
            body_style = ParagraphStyle(
                "DocPilotBody", parent=styles["BodyText"], fontName=font_name,
                fontSize=10, leading=16, textColor="#344c55", spaceAfter=8,
            )
            stream = BytesIO()
            document = SimpleDocTemplate(
                stream, pagesize=A4, leftMargin=22 * mm, rightMargin=22 * mm,
                topMargin=24 * mm, bottomMargin=22 * mm,
                title=title, author="DocPilot",
            )
            story: list[Any] = [Spacer(1, 35 * mm), Paragraph(html.escape(title), title_style)]
            story.extend(Paragraph(html.escape(line), meta_style) for line in lines)
            story.extend([PageBreak(), Paragraph("Текст документа", title_style)])
            for section in sections:
                story.append(Paragraph(html.escape(f"{section['no']} {section['title']}"), heading_style))
                for paragraph in section["text"].splitlines():
                    story.append(Paragraph(html.escape(paragraph), body_style))
            document.build(story)
            return stream.getvalue()
        except (ImportError, OSError, ValueError):
            return cls._minimal_pdf(title, [*lines, *[f"{section['no']} {section['title']}" for section in sections]])

    @staticmethod
    def _authors(source: dict[str, Any], reviewer: str | None, released_by: str) -> list[str]:
        ordered = [element["author"] for section in source["sections"] for element in section["elements"]]
        ordered.extend([reviewer, released_by])
        return list(dict.fromkeys(author for author in ordered if author))

    def create_release(self, rendering_id: str, released_by: str = "si", include_docx: bool = False) -> dict[str, Any]:
        rendering = self._rendering(rendering_id)
        if not rendering["accepted_at"]:
            raise DocPilotError(409, "RENDERING_NOT_ACCEPTED", "Сначала примите рендеринг")
        if not rendering["baseline_id"]:
            raise DocPilotError(409, "RELEASE_BASELINE_REQUIRED", "Выпуск возможен только из базированной структуры")
        baseline = self._baseline(rendering["baseline_id"]) if rendering["baseline_id"] else None
        source = baseline["snapshot"] if baseline else self._snapshot(rendering["document_id"])
        authors = self._authors(source, rendering["reviewer"], released_by)
        document = self._snapshot(rendering["document_id"])
        native_print = os.getenv("DOCPILOT_NATIVE_PRINT") == "1"
        strictdoc = find_tool("strictdoc")
        typst = find_tool("typst")
        template = Path(document["git_path"]) / "docs" / "typst" / "conops.typ"
        if native_print and not strictdoc:
            raise DocPilotError(
                503, "PRINT_HTML2PDF_UNAVAILABLE", "StrictDoc HTML2PDF недоступен",
                ["Исполняемый файл strictdoc не найден."],
            )
        if native_print and (not typst or not template.exists()):
            reason = "Исполняемый файл typst не найден." if not typst else f"Шаблон не найден: {template}"
            raise DocPilotError(503, "PRINT_TYPST_UNAVAILABLE", "Typst недоступен", [reason])
        release_id = f"release-{uuid.uuid4().hex[:10]}"
        release_dir = Path(document["git_path"]) / "releases" / release_id
        release_dir.mkdir(parents=True, exist_ok=True)

        def native_failure(status: int, code: str, message: str, reasons: list[str]) -> DocPilotError:
            shutil.rmtree(release_dir, ignore_errors=True)
            return DocPilotError(status, code, message, reasons)

        baseline_label = baseline["git_tag"] if baseline else "PREVIEW"
        html_path = release_dir / "ConOps.html"
        html_body = "".join(
            f"<section><h2>{html.escape(section['no'])} {html.escape(section['title'])}</h2><p>{html.escape(section['text']).replace(chr(10), '<br>')}</p></section>"
            for section in rendering["sections"]
        )
        html_path.write_text(
            f"<!doctype html><meta charset='utf-8'><title>{html.escape(document['title'])}</title>"
            f"<style>body{{font:15px sans-serif;max-width:780px;margin:60px auto;color:#17303a}}h1{{font-size:28px}}"
            f"code{{color:#176b87}}section{{margin-top:32px}}</style><h1>{html.escape(document['title'])}</h1>"
            f"<p>Проект: {html.escape(document['project_name'])}<br>Базирование: <code>{baseline_label}</code>"
            f"<br>Авторы: {html.escape(', '.join(authors))}</p>{html_body}",
            encoding="utf-8",
        )
        summary = [f"Project: {document['project_name']}", f"Baseline: {baseline_label}", f"Authors: {', '.join(authors)}", f"Rendering: v{rendering['version']}"]
        html_pdf = release_dir / "ConOps-html2pdf.pdf"
        typst_pdf = release_dir / "ConOps-typst.pdf"
        html2pdf_mode = "fallback"
        rendered_sdoc = release_dir / f"{document['code']}-rendered.sdoc"
        release_title = (
            f"{document['project_name']} · {document['title']} · {baseline_label} · "
            f"Авторы: {', '.join(authors)}"
        )
        rendered_sdoc.write_text(
            generate_rendered_sdoc(release_title, rendering, self._grammar("rendered.sgra")),
            encoding="utf-8",
        )
        if native_print:
            assert strictdoc is not None
            strictdoc_output = release_dir / "strictdoc-output"
            strictdoc_config = release_dir / "strictdoc.toml"
            strictdoc_config.write_text(
                '[project]\ntitle = "DocPilot release"\nfeatures = ["HTML2PDF"]\n',
                encoding="utf-8",
            )
            try:
                command = [
                    strictdoc,
                    "export",
                    str(rendered_sdoc),
                    "--formats",
                    "html2pdf",
                    "--output-dir",
                    str(strictdoc_output),
                    "--no-parallelization",
                    "--config",
                    str(strictdoc_config),
                ]
                chromedriver = find_tool("chromedriver")
                if chromedriver:
                    command.extend(["--chromedriver", chromedriver])
                result = subprocess.run(
                    command,
                    text=True,
                    capture_output=True,
                    check=False,
                    timeout=120,
                )
                generated_pdf = next(strictdoc_output.rglob("*.pdf"), None) if result.returncode == 0 else None
            except subprocess.TimeoutExpired as error:
                raise native_failure(
                    502, "PRINT_HTML2PDF_FAILED", "StrictDoc HTML2PDF не завершил выпуск",
                    [f"Превышен тайм-аут {error.timeout} с."],
                ) from error
            except OSError as error:
                raise native_failure(
                    502, "PRINT_HTML2PDF_FAILED", "StrictDoc HTML2PDF не запущен", [str(error)],
                ) from error
            if not generated_pdf:
                diagnostic = (result.stderr or result.stdout or "PDF не создан").strip()[-2000:]
                raise native_failure(
                    502, "PRINT_HTML2PDF_FAILED", "StrictDoc HTML2PDF завершился с ошибкой", [diagnostic],
                )
            shutil.copy2(generated_pdf, html_pdf)
            html2pdf_mode = "strictdoc-html2pdf"
        else:
            html_pdf.write_bytes(self._fallback_pdf("DocPilot ConOps / HTML2PDF fallback", summary, rendering["sections"]))
        typst_mode = "fallback"
        if typst and template.exists():
            body = "\n\n".join(f"{section['no']} {section['title']}\n{section['text']}" for section in rendering["sections"])
            try:
                result = subprocess.run(
                    [
                        typst, "compile", str(template), str(typst_pdf),
                        "--input", f"project={document['project_name']}",
                        "--input", f"title={document['title']}",
                        "--input", f"baseline={baseline_label}",
                        "--input", f"authors={', '.join(authors)}",
                        "--input", f"body={body}",
                    ],
                    text=True, capture_output=True, check=False, timeout=120,
                )
                if result.returncode == 0 and typst_pdf.exists():
                    typst_mode = "typst"
                elif native_print:
                    diagnostic = (result.stderr or result.stdout or "PDF не создан").strip()[-2000:]
                    raise native_failure(
                        502, "PRINT_TYPST_FAILED", "Typst завершился с ошибкой", [diagnostic],
                    )
                else:
                    typst_pdf.write_bytes(self._fallback_pdf("DocPilot ConOps / Typst fallback", summary, rendering["sections"]))
            except subprocess.TimeoutExpired as error:
                if native_print:
                    raise native_failure(
                        502, "PRINT_TYPST_FAILED", "Typst не завершил выпуск",
                        [f"Превышен тайм-аут {error.timeout} с."],
                    ) from error
                typst_pdf.write_bytes(self._fallback_pdf("DocPilot ConOps / Typst fallback", summary, rendering["sections"]))
            except OSError as error:
                if native_print:
                    raise native_failure(
                        502, "PRINT_TYPST_FAILED", "Typst не запущен", [str(error)],
                    ) from error
                typst_pdf.write_bytes(self._fallback_pdf("DocPilot ConOps / Typst fallback", summary, rendering["sections"]))
        else:
            if native_print:
                reason = "Исполняемый файл typst не найден." if not typst else f"Шаблон не найден: {template}"
                raise native_failure(503, "PRINT_TYPST_UNAVAILABLE", "Typst недоступен", [reason])
            typst_pdf.write_bytes(self._fallback_pdf("DocPilot ConOps / Typst fallback", summary, rendering["sections"]))
        files: dict[str, Any] = {
            "html": {"path": str(html_path), "mode": "native"},
            "pdf_html2pdf": {"path": str(html_pdf), "mode": html2pdf_mode},
            "pdf_typst": {"path": str(typst_pdf), "mode": typst_mode},
        }
        if include_docx:
            docx_path = release_dir / "ConOps.docx"
            try:
                from docx import Document

                docx = Document()
                docx.add_heading(document["title"], 0)
                docx.add_paragraph(f"Базирование: {baseline_label}")
                docx.add_paragraph(f"Авторы: {', '.join(authors)}")
                for section in rendering["sections"]:
                    docx.add_heading(f"{section['no']} {section['title']}", 1)
                    docx.add_paragraph(section["text"])
                docx.save(docx_path)
                files["docx"] = {"path": str(docx_path), "mode": "python-docx"}
            except ImportError:
                files["docx"] = {"path": None, "mode": "unavailable"}
        at = utcnow()
        with self.database.connect() as connection:
            connection.execute(
                "INSERT INTO release(id, rendering_id, files, authors, released_by, at) VALUES (?, ?, ?, ?, ?, ?)",
                (release_id, rendering_id, json.dumps(files, ensure_ascii=False), json.dumps(authors, ensure_ascii=False), released_by, at),
            )
        logger.info("release created id=%s rendering=%s formats=%s", release_id, rendering_id, sorted(files))
        return {"id": release_id, "rendering_id": rendering_id, "files": files, "authors": authors, "released_by": released_by, "at": at}

    def list_releases(self, document_id: str) -> list[dict[str, Any]]:
        with self.database.connect() as connection:
            rows = connection.execute(
                """SELECT rel.* FROM release rel JOIN rendering r ON r.id=rel.rendering_id
                   WHERE r.document_id=? ORDER BY rel.at DESC, rel.rowid DESC""", (document_id,)
            ).fetchall()
        return [
            {
                "id": row["id"], "rendering_id": row["rendering_id"], "files": _json(row["files"], {}),
                "authors": _json(row["authors"], []), "released_by": row["released_by"], "at": row["at"],
            }
            for row in rows
        ]

    def release_file(self, release_id: str, file_key: str) -> Path:
        with self.database.connect() as connection:
            row = connection.execute("SELECT files FROM release WHERE id = ?", (release_id,)).fetchone()
        if not row:
            raise DocPilotError(404, "RELEASE_NOT_FOUND", "Выпуск не найден", [release_id])
        entry = _json(row["files"], {}).get(file_key)
        if not entry or not entry.get("path"):
            raise DocPilotError(404, "RELEASE_FILE_NOT_FOUND", "Файл выпуска не найден", [file_key])
        path = Path(entry["path"]).resolve()
        if not path.is_file() or "releases" not in path.parts:
            raise DocPilotError(404, "RELEASE_FILE_NOT_FOUND", "Файл выпуска недоступен", [file_key])
        return path
